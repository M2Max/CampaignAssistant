import json
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, List, Tuple

import fitz  # PyMuPDF
import yaml
from docx import Document as DocxDocument
from langchain_core.documents import Document

SUPPORTED_FILE_TYPES: List[str] = ["txt", "md", "json", "yaml", "yml", "pdf", "docx"]


def _docx_to_markdown(file_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(file_bytes))
    markdown_lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if "heading" in style_name:
            level = 1
            for num in range(1, 7):
                if f"heading {num}" in style_name:
                    level = num
                    break
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            markdown_lines.append(text)
    return "\n\n".join(markdown_lines)


def _pdf_to_text(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = [page.get_text("text") for page in doc]
    return "\n\n".join(pages)


def extract_text_from_upload(uploaded_file) -> List[str]:
    """Normalise supported file uploads into plain-text snippets."""
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    if suffix in {".txt", ".md"}:
        return [file_bytes.decode("utf-8", errors="ignore")]

    if suffix == ".json":
        parsed = json.loads(file_bytes.decode("utf-8", errors="ignore"))
        return [json.dumps(parsed, indent=2)]

    if suffix in {".yaml", ".yml"}:
        parsed = yaml.safe_load(file_bytes.decode("utf-8", errors="ignore"))
        return [yaml.safe_dump(parsed, sort_keys=False)]

    if suffix == ".pdf":
        return [_pdf_to_text(file_bytes)]

    if suffix == ".docx":
        return [_docx_to_markdown(file_bytes)]

    raise ValueError(f"Unsupported file type: {suffix}")


def ingest_documents(uploaded_files: Iterable, campaign_history: Any) -> Tuple[str | None, str | None]:
    """Process uploaded files and push resulting documents into the vector store."""
    documents: List[Document] = []
    errors: List[str] = []

    for uploaded_file in uploaded_files:
        try:
            texts = extract_text_from_upload(uploaded_file)
            for idx, text in enumerate(texts):
                cleaned = text.strip()
                if not cleaned:
                    continue
                documents.append(
                    Document(
                        page_content=cleaned,
                        metadata={
                            "source": uploaded_file.name,
                            "chunk": idx,
                        },
                    )
                )
        except Exception as exc:
            errors.append(f"{uploaded_file.name}: {exc}")

    if documents:
        campaign_history.add_documents(documents)

    success_msg = f"Added {len(documents)} document chunks to the vectorstore." if documents else None
    error_msg = "\n".join(errors) if errors else None
    return success_msg, error_msg
