import json
import time
from io import BytesIO
from pathlib import Path
from typing import List, Optional
from uuid import uuid4

import streamlit as st
import yaml
from dotenv import load_dotenv
from langchain_core.documents import Document
from docx import Document as DocxDocument
import fitz  # PyMuPDF

from Agent.CampaignAgent import CampaignAgent

load_dotenv()

st.set_page_config(
    page_title="Campaign Assistant",
    page_icon="📊",
    layout="wide",
)

SUPPORTED_FILE_TYPES = ["txt", "md", "json", "yaml", "yml", "pdf", "docx"]


def inject_minimal_theme():
    st.markdown(
        """
        <style>
        :root {
            color-scheme: light dark;
        }
        .stApp {
            background-color: var(--background-color);
            color: var(--text-color);
        }
        section[data-testid="stSidebar"] {
            background: var(--secondary-background-color);
            background: color-mix(in srgb, var(--secondary-background-color) 90%, var(--background-color) 10%);
            color: var(--text-color);
            border-right: 1px solid color-mix(in srgb, var(--text-color) 8%, transparent);
        }
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] label,
        section[data-testid="stSidebar"] .stMarkdown {
            color: var(--text-color);
        }
        .stButton>button {
            background: var(--secondary-background-color);
            background: color-mix(in srgb, var(--primary-color) 12%, transparent);
            color: var(--text-color);
            border-radius: 10px;
            border: 1px solid color-mix(in srgb, var(--primary-color) 30%, transparent);
            backdrop-filter: blur(6px);
        }
        .stButton>button:hover {
            border-color: var(--primary-color);
        }
        .minimal-card {
            background: var(--secondary-background-color);
            background: color-mix(in srgb, var(--secondary-background-color) 94%, var(--background-color) 6%);
            color: var(--text-color);
            padding: 1.4rem;
            border-radius: 18px;
            border: 1px solid color-mix(in srgb, var(--text-color) 10%, transparent);
            box-shadow: 0 12px 28px color-mix(in srgb, var(--text-color) 6%, transparent);
        }
        div[data-testid="stChatMessageContent"] {
            background: var(--secondary-background-color);
            background: color-mix(in srgb, var(--secondary-background-color) 92%, var(--background-color) 8%);
            border-radius: 16px;
            border: 1px solid color-mix(in srgb, var(--text-color) 8%, transparent);
            padding: 1rem 1.25rem;
        }
        div[data-testid="stChatMessageContent"] p {
            color: var(--text-color);
        }
        section[data-testid="stSidebar"] .stFileUploader label {
            color: var(--text-color);
        }
        section[data-testid="stSidebar"] .stAlert {
            background: color-mix(in srgb, var(--secondary-background-color) 90%, var(--background-color) 10%);
            border-radius: 12px;
        }
        [data-baseweb="input"] input,
        [data-baseweb="textarea"] textarea,
        .stTextInput input,
        .stTextArea textarea {
            background: var(--secondary-background-color);
            background: color-mix(in srgb, var(--secondary-background-color) 92%, var(--background-color) 8%);
            color: var(--text-color);
            border-radius: 12px;
        }
        .stTextInput input:focus,
        .stTextArea textarea:focus {
            border-color: color-mix(in srgb, var(--primary-color) 40%, transparent);
            box-shadow: 0 0 0 1px color-mix(in srgb, var(--primary-color) 25%, transparent);
        }
        .stMarkdown p,
        .stMarkdown li {
            color: var(--text-color);
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def ensure_session_state():
    if "chat_sessions" not in st.session_state:
        st.session_state.chat_sessions = []
    if "active_chat_id" not in st.session_state:
        st.session_state.active_chat_id = None
    if "show_uploader" not in st.session_state:
        st.session_state.show_uploader = False
    if "ingestion_feedback" not in st.session_state:
        st.session_state.ingestion_feedback = None
    if "vectorstore_uploader_key" not in st.session_state:
        st.session_state.vectorstore_uploader_key = "vectorstore_uploader"
    if not st.session_state.chat_sessions:
        create_new_chat()


def create_new_chat() -> dict:
    chat_id = str(uuid4())
    new_chat = {
        "id": chat_id,
        "title": f"Untitled chat {len(st.session_state.chat_sessions) + 1}",
        "messages": [],
        "created_at": time.time(),
    }
    st.session_state.chat_sessions.insert(0, new_chat)
    st.session_state.active_chat_id = chat_id
    return new_chat


def get_active_chat() -> Optional[dict]:
    for chat in st.session_state.chat_sessions:
        if chat["id"] == st.session_state.active_chat_id:
            return chat
    return None


def rename_chat_if_needed(chat: dict, prompt: str):
    if chat["title"].startswith("Untitled"):
        trimmed = prompt.strip().splitlines()[0]
        if trimmed:
            chat["title"] = (trimmed[:32] + "…") if len(trimmed) > 32 else trimmed


def handle_chat_selection():
    if not st.session_state.chat_sessions:
        return

    active_index = 0
    for idx, chat in enumerate(st.session_state.chat_sessions):
        if chat["id"] == st.session_state.active_chat_id:
            active_index = idx
            break

    selection = st.radio(
        "Previous chats",
        options=list(range(len(st.session_state.chat_sessions))),
        index=active_index,
        format_func=lambda idx: st.session_state.chat_sessions[idx]["title"],
        label_visibility="collapsed",
    )
    st.session_state.active_chat_id = st.session_state.chat_sessions[selection]["id"]


def docx_to_markdown(file_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(file_bytes))
    markdown_lines = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style and para.style.name else ""
        if "heading" in style:
            level = 1
            for num in range(1, 7):
                if f"heading {num}" in style:
                    level = num
                    break
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            markdown_lines.append(text)
    return "\n\n".join(markdown_lines)


def pdf_to_text(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = [page.get_text("text") for page in doc]
    return "\n\n".join(pages)


def extract_text_from_upload(uploaded_file) -> List[str]:
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)

    if suffix in {".txt", ".md"}:
        return [file_bytes.decode("utf-8", errors="ignore")]

    if suffix == ".json":
        parsed = json.loads(file_bytes.decode("utf-8", errors="ignore"))
        return [json.dumps(parsed, indent=2)]

    if suffix in {".yaml", ".yml"}:
        parsed = yaml.safe_load(file_bytes.decode("utf-8", errors="ignore"))
        return [yaml.safe_dump(parsed, sort_keys=False)]

    if suffix == ".pdf":
        return [pdf_to_text(file_bytes)]

    if suffix == ".docx":
        return [docx_to_markdown(file_bytes)]

    raise ValueError(f"Unsupported file type: {suffix}")


def ingest_documents(uploaded_files, agent: CampaignAgent):
    documents = []
    errors = []

    for uploaded_file in uploaded_files:
        try:
            texts = extract_text_from_upload(uploaded_file)
            for idx, text in enumerate(texts):
                cleaned = text.strip()
                if not cleaned:
                    continue
                documents.append(
                    Document(
                        page_content=cleaned,
                        metadata={
                            "source": uploaded_file.name,
                            "chunk": idx,
                        },
                    )
                )
        except Exception as exc:
            errors.append(f"{uploaded_file.name}: {exc}")

    if documents:
        agent.campaign_history.add_documents(documents)

    success_msg = f"Added {len(documents)} document chunks to the vectorstore." if documents else None
    error_msg = "\n".join(errors) if errors else None
    return success_msg, error_msg


@st.cache_resource
def get_agent():
    return CampaignAgent()


def render_sidebar(agent: CampaignAgent):
    with st.sidebar:
        st.markdown("### Workspace")
        new_chat_clicked = st.button("＋ New Chat", use_container_width=True)
        add_docs_clicked = st.button("⬆️ Add documents to vectorstore", use_container_width=True)

        if new_chat_clicked:
            create_new_chat()
            st.rerun()

        if add_docs_clicked:
            st.session_state.show_uploader = not st.session_state.show_uploader

        if st.session_state.show_uploader:
            st.write("Upload .txt, .md, .json, .yaml/.yml, .pdf, .docx")
            uploaded_files = st.file_uploader(
                "Select documents",
                type=SUPPORTED_FILE_TYPES,
                accept_multiple_files=True,
                key=st.session_state.vectorstore_uploader_key,
            )
            if uploaded_files:
                success, error = ingest_documents(uploaded_files, agent)
                st.session_state.ingestion_feedback = {"success": success, "error": error}
                st.session_state.show_uploader = False
                st.session_state.vectorstore_uploader_key = f"vectorstore_uploader_{uuid4().hex}"
                st.rerun()

        if st.session_state.ingestion_feedback:
            feedback = st.session_state.ingestion_feedback
            if feedback.get("success"):
                st.success(feedback["success"])
            if feedback.get("error"):
                st.warning(feedback["error"])

        st.markdown("---")
        st.markdown("### Past chats")
        handle_chat_selection()


def render_chat_ui(agent: CampaignAgent):
    active_chat = get_active_chat()
    if not active_chat:
        active_chat = create_new_chat()

    st.title("📊 Campaign Assistant")
    st.markdown(
        '<div class="minimal-card">Ask anything about your marketing campaign history. '
        "Responses will use the current vectorstore context.</div>",
        unsafe_allow_html=True,
    )

    for message in active_chat["messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Ask about your campaigns…"):
        active_chat["messages"].append({"role": "user", "content": prompt})
        rename_chat_if_needed(active_chat, prompt)

        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            final_response = ""
            response_placeholder = st.empty()

            try:
                with st.spinner("Thinking through your campaign data…"):
                    result = agent.invoke(prompt)
                    final_response = result["messages"][-1].content

                response_placeholder.markdown(final_response)

            except Exception as e:
                final_response = f"An error occurred: {str(e)}"
                response_placeholder.error(final_response)
                print(f"Error: {e}")

        active_chat["messages"].append(
            {
                "role": "assistant",
                "content": final_response if final_response else "No response generated",
            }
        )


def main():
    inject_minimal_theme()
    ensure_session_state()
    agent = get_agent()
    render_sidebar(agent)
    render_chat_ui(agent)


if __name__ == "__main__":
    main()